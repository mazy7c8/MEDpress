from subprocess import CalledProcessError
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.font as tf
from tkinter import Canvas, Event, Frame, INSERT, END, PhotoImage, filedialog

from jinja2.loaders import BaseLoader
from createtmp import template_window
from item import ListItem, readFolder, readTemplate, readBody
from jinja2 import Template, Environment, FileSystemLoader, select_autoescape, meta
from jinja2schema import infer, model, config
import subprocess
import re
import time
import os
from draw import vardrawing
from idlelib.tooltip import Hovertip
import ntpath
from docx import Document
from apscheduler.schedulers.background import BackgroundScheduler

class MEDpress(object):        
    def __init__(self, parent):
        self.root = parent
        self.root.title("Ekran Glowny")
        self.frame = tk.Frame(parent)
        self.frame.pack()
        self.root['bg'] = "#FCAFAF"
        self.root.protocol("WM_DELETE_WINDOW", self.root.quit())

        def_font = tf.nametofont("TkDefaultFont")
        def_font.configure(family='Georgia')
        self.root.option_add("*Font",def_font)

        self.entryBoxList = {}
        self.found = ListItem

        self.templateStack = []
        self.fromFileStack = []

        self.dummyTemplate = ListItem('','','')

        self.actualDrawings = []

        datafromfolder = readFolder()
        for source, name, time in datafromfolder:
            testowe = ListItem(name, time, source)
            self.templateStack.append(testowe)

        self.JinjaEnv = Environment(
            loader=FileSystemLoader('szablony'),
            autoescape=select_autoescape(['txt', 'xml'])
        )

        #template = self.JinjaEnv.get_template('szablon.txt')
        #druk = template.render(imie="Jan",nazwisko="Kowalski")

        canvas = Canvas(self.frame, width=1300, height=900, bg="#FCAFAF")
        canvas.create_rectangle(492, 260, 1287, 884, fill="#ED6868")
        canvas.pack()

        def focus_next_widget(event):
            event.widget.tk_focusNext().focus()
            return("break")

        button1 = tk.Button(
            self.frame,
            text="Stwórz szablon",
            width=20,
            height=10,
            bg="#DCE19C",
            command=self.openFrame,
            takefocus=0
        )
        button1.pack()
        button1.place(y=44, x=79)
        myTip1 = Hovertip(button1,"Tutaj zacznij tworzenie nowego formularza od podstaw")

        button2 = tk.Button(
            self.frame,
            text="Wgraj z pliku",
            width=20,
            height=10,
            bg="#C0D9B7",
            takefocus=0,
            command=self.openFromFile,

        )
        button2.pack()
        button2.place(y=44, x=279)
        myTip2 = Hovertip(button2,"Wgraj formularz znajdujacy sie poza folderem szablonow na liste")

        button3 = tk.Button(
            self.frame,
            text="Chmura szablonów",
            width=20,
            height=10,
            bg="#B2D4DC",
            takefocus=0
        )
        button3.pack()
        button3.place(y=44, x=479)
        myTip3 = Hovertip(button3,"Świat formularzy z internetu")


        self.textfield = tk.Text(
            self.frame,
            bg='white',
        )
        #textfield.insert(INSERT, druk)
        self.textfield.pack()
        self.textfield.place(y=44, x=714, height=147, width=573)
        self.textfield.bind("<Tab>", focus_next_widget)
        self.textfield.bind("<Return>", lambda event : self.getTextEntry())
        myTip4 = Hovertip(self.textfield,"Aby zaczac wpisz skrot formularza z listy po lewej")



        texfieldlabel = tk.Label(
            self.frame,
            text="Wpisz skroty szablonow/gotowy tekst",
            font=("Helvetica", 10),
            bg=self.root['bg']
        )
        texfieldlabel.pack()
        texfieldlabel.place(x=713, y=17, height=20, width=379)

        keyinfolabel = tk.Label(
            self.frame,
            text="Ctrl+s rozpocznij wypis | Ctrl+z zakoncz wypis | Ctrl+c do schowka",
            font=("Helvetica", 10),
            bg=self.root['bg']
        )
        keyinfolabel.pack()
        keyinfolabel.place(x=100, y=17, height=20, width=500)
        myTip5 = Hovertip(keyinfolabel,"Najwazniejsze skroty klawiszowe w programie")


        Xbutton = tk.Button(
            self.frame,
            text="X",
            width=4,
            height=2,
            bg="lightgrey",
            command=self.cleanTextfield,
            takefocus=0
        )
        Xbutton.pack()
        Xbutton.place(x=1250, y=7)
        myTip6 = Hovertip(Xbutton,"Wyczysc pole")

        self.tree = ttk.Treeview(
            self.frame,
            takefocus=1
            )
        self.tree["columns"] = ("COL2", "COL3")
        self.tree.column("#0", width=100)
        self.tree.heading("#0", text="Nazwa", anchor=tk.W)
        self.tree.column("COL2", width=100)
        self.tree.heading("COL2", text="Skrót", anchor=tk.W)
        self.tree.column("COL3", width=100)
        self.tree.heading("COL3", text="Data edycji", anchor=tk.W)
        number = 0
        for testowe in self.templateStack:
            number += 1
            self.tree.insert('', 'end', 'ID{0}'.format(
                number), text=testowe.name, values=(testowe.abbr, testowe.date))
        self.tree.pack()
        self.tree.place(x=44, y=251, height=633, width=382)
        self.tree.bind('<ButtonRelease-1>', self.updateTextfieldFromClick)
        self.tree.bind('<Return>', self.updateTextfieldFromClick)
        try:
            child_id = self.tree.get_children()[0]
        except IndexError:
            self.tree.insert('', 'end', 'ID{0}'.format(number), text="Brak szablonow", values=("w folderze szablony", "00:00:00"))
            child_id = self.tree.get_children()[0]
        self.tree.focus_set()
        self.tree.focus(child_id)
        self.tree.selection_set(child_id)
        #myTip7 = Hovertip(tk.Widget,"Tutaj znajduja sie wszystkie formularze z katalogu szablony")


        treelabel = tk.Label(
            self.frame,
            text="Twoje szablony",
            font=("Helvetica", 10),
            bg=self.root['bg']
        )
        treelabel.pack()
        treelabel.place(x=20, y=222, height=20, width=150)

        editbutton = tk.Button(
            self.frame,
            text="Edytuj w edytorze",
            width=14,
            height=1,
            bg="#DCE19C",
            command=self.openFrameWithTmp,
            takefocus=0
        )
        editbutton.pack()
        editbutton.place(y=224, x=320)
        myTip8 = Hovertip(editbutton,"Użyj aby podejrzec zawartosc formularza lub wprowadzic w nim zmiany")


        savebutton = tk.Button(
            self.frame,
            text="Zapisz do pliku",
            width=14,
            height=1,
            bg="#C0D9B7",
            takefocus=0
        )
        savebutton.pack()
        savebutton.place(y=224, x=460)
        myTip9 = Hovertip(savebutton,"Eksportuj wszystkie formularze do archiwum")

        refreshbutton = tk.Button(
            self.frame,
            text="Odswiez",
            width=14,
            height=1,
            bg="#00ff00",
            command=self.refreshTmpList,
            takefocus=0
        )
        refreshbutton.pack()
        refreshbutton.place(y=224, x=180)
        myTip10 = Hovertip(refreshbutton,"Odśwież listę formularzy i wyczyść okno rysowania")

        self.startbutton = tk.Button(
            self.frame,
            text="Rozpocznij wypis",
            width=14,
            height=2,
            bg="lightgrey",
            activebackground='#00ff00',
            command=self.getTextEntry,
            takefocus=1

        )
        self.startbutton.pack()
        self.startbutton.place(x=840, y=201)
        self.startbutton.bind("<Return>", lambda event : self.getTextEntry())
        myTip11 = Hovertip(self.startbutton,"Zacznij rysowanie formularza z pola powyzej")


        self.clipboardbutton = tk.Button(
            self.frame,
            text="Do schowka",
            width=14,
            height=2,
            bg="lightgrey",
            activebackground='#00ff00',
            command=self.copyToClipboard,
            takefocus=0
        )
        self.clipboardbutton.pack()
        self.clipboardbutton.place(x=998, y=201)
        myTip12 = Hovertip(self.clipboardbutton,"Skopiuj pole powyzej do systemowego schowka")

        exportbutton = tk.Button(
            self.frame,
            text="Eksport pdf/rtf",
            width=14,
            height=2,
            bg="lightgrey",
            takefocus=0,
            command=self.printPDF
        )
        exportbutton.pack()
        exportbutton.place(x=1156, y=201)        
        myTip13 = Hovertip(exportbutton,"Otworz powyższe w formie dokumentu tekstowego")


        progresslabel = tk.Label(
            self.frame,
            text="Pasek postępu",
            font=("Helvetica",12),
            bg=("#ED6868")
        )
        progresslabel.pack()
        progresslabel.place(x=1020, y=850)

        self.endworkbutton = tk.Button(
            self.frame,
            text="Zakończ",
            width=10,
            height=2,
            bg="lightgrey",
            activebackground='#00ff00',
            command=self.readWork,
            takefocus=1
        )
        self.endworkbutton.pack()
        self.endworkbutton.place(x=1175, y=839)        
        self.endworkbutton.bind("<Return>", lambda event : self.readWork())
        myTip14 = Hovertip(self.endworkbutton,"Wygeneruj wydruk formularza z podanych zmiennych")


        self.root.bind("<Control-s>", lambda event, : self.getTextEntry())
        self.root.bind("<Control-z>", lambda event, : self.readWork())
        self.root.bind("<Control-c>", lambda event, : self.copyToClipboard())

    def makeProgressBar(self,X):
        self.progressbar = ttk.Progressbar(
        self.frame, orient="horizontal", length=500, mode="determinate", maximum=X*10, value=0)
        self.progressbar.pack()
        self.progressbar.place(x=511, y=853)
        #self.progressAction()
        self.scheduler = BackgroundScheduler()
        self.scheduler.start()
        self.scheduler.add_job(self.progressAction, 'interval', seconds = 2)
        
    # def progressAction(self):
    #     for item in self.actualDrawings:
    #         if(type(item))!=list:
    #             item.bind("<Enter>", lambda event : self.progressClick())

    # def progressClick(self):
    #     self.progressbar['value']+=5

    def progressAction(self):
        readed = {}
        self.progressbar['value']=0
        for keys,values in self.entryBoxList.items():
            try:
                if type(values) is list:
                    final=""
                    for value in values:
                        if value.get()!='':
                            final+=value.get()+", "
                        readed[keys]=final.rstrip(", ")

                else: readed[keys] = self.entryBoxList[keys].get()

                if readed[keys]!="":
                    self.progressbar['value']+=10
            except:
                pass
        

            


    def copyToClipboard(self):
        self.clipboardbutton.config(bg="#00FF00")
        root.after(100, lambda: self.clipboardbutton.config(bg='lightgrey'))
        textEntry = self.textfield.get("1.0", "end-1c")
        try:
            # MACOS
            cmd = 'echo '+textEntry.strip()+'|pbcopy'
            return subprocess.check_call(cmd, shell=True)
        except:             
            print("no macos")
        
        try:
            # WINDOWS
            cmd = 'echo '+textEntry.strip()+'|clip'
            return subprocess.check_call(cmd, shell=True)
        except:
            print("no windows")
        
        try:
            # LINUX
            cmd = 'echo ' + textEntry.strip() + ' | tr -d \'\\n\''
            cmd = cmd + ' | xsel -i --clipboard'
            return subprocess.check_call(cmd, shell=True)
        except:
            print("no linux")

    def getTextEntry(self):
        self.startbutton.config(bg="#00FF00")
        root.after(100, lambda: self.startbutton.config(bg='lightgrey'))
        textEntry = self.textfield.get("1.0", "end-1c")
        textEntry= textEntry.rstrip()
        self.templateSearch(textEntry)

    def refreshTmpList(self, added=None):
        self.tree.delete(*self.tree.get_children())

        datafromfolder = readFolder()

        #print(self.templateStack, self.found)
        # result = zip(datafromfolder,self.templateStack)
        # for data, instance in list(result):
        #     instance.updateInstance(data[1], data[2], data[0])
        #     #self.templateStack.append(nowe)
        
        self.templateStack.clear()
        
        for source, name, time in datafromfolder:
            testowe = ListItem(name, time, source)
            self.templateStack.append(testowe)

        if added:
            name, source, time = added
            testowe = ListItem(name, time, source)
            testowe.updateInstance(name, time, source)
            self.fromFileStack.append(testowe)

        self.templateStack.extend(self.fromFileStack)

        
        number = 0
        for cos in self.templateStack:
            number += 1
            self.tree.insert('', 'end', 'ID{0}'.format(
                number), text=cos.name, values=(cos.abbr, cos.date))
        
        self.clearWorkArea()
        self.tree.focus_set()
        self.tree.focus(self.tree.get_children()[0])
        self.tree.selection_set(self.tree.get_children()[0])


    def clearWorkArea(self):
        for drawing in self.actualDrawings:
            if type(drawing)==list:
                for draw in drawing:
                    try: 
                        draw.destroy()
                    except:
                        pass
            try: 
                drawing.destroy()
            except:
                pass
        self.actualDrawings.clear()
            
    def templateSearch(self, string):
        self.clearWorkArea()
        for template in self.templateStack:
            if string == template.abbr:
                self.found = template
                foundvars = self.getVariablesFromTemp(self.found)
                self.entryBoxList = self.drawRequests(foundvars,template)
                self.endworkbutton.lift()
                self.makeProgressBar(len(foundvars))
                break

    def initializeRender(self, object, dictonary):
        template = self.JinjaEnv.get_template(object.source)
        self.druk = template.render(dictonary)
        self.updateTextfield()

    def updateTextfield(self):
        self.textfield.delete('1.0', END)
        try:
            stripped = re.search(r'\n\n.*',self.druk)
            stripped = re.sub(r'^$\n', '', stripped.group(0), flags=re.MULTILINE)
            self.textfield.insert(INSERT, stripped)
        except AttributeError:
            self.textfield.insert(INSERT,self.druk)

    def cleanTextfield(self):
        self.textfield.delete("1.0", END)

    def updateTextfieldFromClick(self, event):
        selected = self.tree.item(self.tree.selection())['values'][0]
        self.textfield.delete('1.0', END)
        self.textfield.insert(INSERT, selected)

    def hide(self):
        self.root.withdraw()

    def openFrame(self):
        self.hide()
        template_window(self, self.dummyTemplate)

    def openFrameWithTmp(self):
        self.hide()
        try:
            selected = self.tree.item(self.tree.selection())['values'][0]
            self.templateSearch(selected)
            template_window(self, self.found)
        except:
            template_window(self, self.dummyTemplate)
        
    def frameHandler(self, otherFrame):
        def handler(): return self.onCloseOtherFrame(otherFrame)
        return handler

    def onCloseOtherFrame(self, otherFrame):
        otherFrame.destroy()
        self.show()

    def saveFromTempEdit(self, template):
        self.tree.insert('', 'end', text=template.name, values=(template.abbr, template.date))  

    def show(self):
        self.root.update()
        self.root.deiconify()
        self.refreshTmpList()
        self.clearWorkArea()

    def getVariablesFromTemp(self, object):
        varlist = []
        try:
            template_source = self.JinjaEnv.loader.get_source(
            self.JinjaEnv, object.source)[0]
            parsed_content = self.JinjaEnv.parse(template_source)
        except:
            template_source = readBody(object,object.name,object.source)

        #varlist = list(meta.find_undeclared_variables(parsed_content))
        
        ordered = []
        for m in  re.finditer(r'\{\{\s*(\w+)\s*\}\}',template_source):
            ordered.append(m.group(1))
        #print(ordered)

        return ordered

    def drawRequests(self, lista, template):

        # vartext = {}
        # varentry = {}

        drawing = {}
        verticalpos = 300
        horizontalpos = 550

        # x = 0
        # for item in lista:
        #     vartext[item] = "texfield{0}".format(x)
        #     varentry[item] = "entrybox{0}".format(x)
        #     x += 1

        for item in lista:
            if verticalpos>=800:
                verticalpos = 300
                horizontalpos = 950

            try:
                drawing[item]=vardrawing(item,self.frame,template.dictionary[item],horizontalpos,verticalpos,template.dictionary)

                verticalpos+=50
                if template.dictionary[item][0]=="CB":
                    verticalpos+=20*len(template.dictionary[item][2:])
                if template.dictionary[item]=="DT" or template.dictionary[item][0]=="DT":
                    verticalpos+=15
                if template.dictionary[item][0]=="TN":
                    verticalpos+=15
                if template.dictionary[item][0]=="NB":
                    verticalpos+=15
                if template.dictionary[item][0]=="IF":
                    passing = template.dictionary[item][2]
                    length = template.dictionary[passing][1:]
                    verticalpos+=20+15*(len(length))
                if template.dictionary[item][0]=="RC":
                    verticalpos+=20*len(template.dictionary[item][2:])
                
            except KeyError:
                drawing[item]=vardrawing(item,self.frame,"TX",horizontalpos,verticalpos,template.dictionary)
                verticalpos+=50


        listofbodies = {}
        for item in lista:
            heading = drawing[item].drawheading()
            self.actualDrawings.append(heading)

            body, bodies = drawing[item].drawbody()
            self.actualDrawings.append(bodies)

            listofbodies[item]=body

        varentry = listofbodies
            
        # for item in lista:
        #     vartext[item] = tk.Label(
        #         self.frame,
        #         text="Zmienna "+str(item)+" typu:"+str(template.dictionary[item]),
        #         font=("Helvetica", 16),
        #         bg=self.root['bg']
        #     )
        #     vartext[item].pack()
        #     vartext[item].place(x=550, y=verticalpos, height=20, width=300)
        #     self.actualDrawings.append(vartext[item])

        #     varentry[item] = tk.Entry(
        #         self.frame,
        #     )
        #     varentry[item].pack()
        #     varentry[item].place(x=550, y=verticalpos+20, height=20, width=300)
        #     varentry[item].focus_set()
        #     self.actualDrawings.append(varentry[item])

        #     verticalpos += 100

        return varentry

    def readWork(self):
        self.scheduler.remove_all_jobs()
        self.endworkbutton.config(bg="#00FF00")
        root.after(100, lambda: self.endworkbutton.config(bg='lightgrey'))
        readed = {}
        #print(self.entryBoxList)
        for keys,values in self.entryBoxList.items():
            try:
                if type(values) is list:
                    final=""
                    for value in values:
                        if value.get()!='':
                            final+=value.get()+", "
                        readed[keys]=final.rstrip(", ")

                else: readed[keys] = self.entryBoxList[keys].get()
            except AttributeError:
                readed[keys]="#$@%*!"
        self.initializeRender(self.found, readed)

    def alternativeRender(self):
        textEntry = self.textfield.get("1.0", "end-1c")
        
        for template in self.templateStack:
            if textEntry == template.abbr:
                self.found = template
        
        data = readTemplate(self.found)

        textlabel = tk.Label(
            self.frame,
            font=("Helvetica", 10),
            bg=['#ED6868']
        )
        textlabel.pack()
        textlabel.place(x=600, y=300)


        counter=1
        def go(counter):
            text=data[:counter]
            textlabel.config(text=text)
            root.after(10, lambda: go(counter+1))
            if "{" in text: stop()

        def stop():
            root.after(1000,lambda: go(counter))

        go(counter)

    def openFromFile(self):
        #root.withdraw()
        file_path = filedialog.askopenfilename()
        if file_path:
            with open(os.path.join(file_path), "r") as src_file:
                txttime = time.ctime(os.path.getctime(src_file.name))
                head, tail = ntpath.split(file_path)
                #return tail or ntpath.basename(head)
                name=tail.rstrip(".txt")
                data = (name, file_path, txttime)

                self.refreshTmpList(data)
    
    def printPDF(self):
        textEntry = self.textfield.get("1.0", "end-1c")
        header = ListItem.returnName(self.found)
        
        document = Document()
        document.add_heading(header, 0)
        document.add_paragraph(textEntry)
        #document.add_page_break()
        document.save('wydruk.docx')



          




        

if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("1300x900")
    root.call('encoding', 'system', 'utf-8')
    try:
        root.iconbitmap('icon.ico')
        root.tk.call('wm', 'iconphoto', root._w, PhotoImage(file='icon.ico'))
    except:
        pass
    try:
        img = tk.Image("photo", file="icon.png")
        root.tk.call('wm', 'iconphoto', root._w, img)
    except:
        pass
   
    app = MEDpress(root)
    root.mainloop()
